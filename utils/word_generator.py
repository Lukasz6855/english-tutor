# -*- coding: utf-8 -*-
"""
Moduł do generowania plików Word z listą słówek
Tworzy dokumenty w formacie zgodnym z przykładem użytkownika
"""

# Importowanie biblioteki do tworzenia plików Word
from docx import Document
from docx.shared import Pt, Inches, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Importowanie biblioteki do operacji na plikach
from io import BytesIO

# Importowanie wyrażeń regularnych do parsowania tekstu
import re


class WordGenerator:
    """
    Klasa do generowania plików Word z listą słówek
    """
    
    def __init__(self):
        """
        Inicjalizacja generatora dokumentów Word
        """
        # Domyślna czcionka i rozmiar
        self.default_font = "Cambria"
        self.default_font_size = Pt(11)
        # Odstęp między hasłami (4 punkty)
        self.spacing_after = Pt(4)
    
    def _set_narrow_margins(self, doc):
        """
        Ustawia wąskie marginesy w dokumencie
        
        Args:
            doc: Dokument Word
        """
        # Pobieranie sekcji dokumentu
        for section in doc.sections:
            # Wąskie marginesy (1.27 cm = 0.5 cala)
            section.top_margin = Cm(1.27)
            section.bottom_margin = Cm(1.27)
            section.left_margin = Cm(1.27)
            section.right_margin = Cm(1.27)
    
    def _set_two_columns(self, doc):
        """
        Ustawia dwie kolumny w dokumencie
        
        Args:
            doc: Dokument Word
        """
        # Pobieranie sekcji dokumentu
        for section in doc.sections:
            # Tworzenie elementu kolumn
            sectPr = section._sectPr
            
            # Usuwanie istniejących ustawień kolumn
            for cols in sectPr.findall(qn('w:cols')):
                sectPr.remove(cols)
            
            # Tworzenie nowego elementu kolumn
            cols = OxmlElement('w:cols')
            # Ustawienie 2 kolumn
            cols.set(qn('w:num'), '2')
            # Odstęp między kolumnami (0.5 cm)
            cols.set(qn('w:space'), '284')  # 284 twips = ~0.5 cm
            
            # Dodanie do sekcji
            sectPr.append(cols)
    
    def _format_paragraph(self, paragraph, is_bold=False):
        """
        Formatuje paragraf z domyślnymi ustawieniami
        
        Args:
            paragraph: Paragraf do sformatowania
            is_bold: Czy tekst ma być pogrubiony
        """
        # Ustawienie czcionki i rozmiaru dla każdego fragmentu
        for run in paragraph.runs:
            run.font.name = self.default_font
            run.font.size = self.default_font_size
            run.bold = is_bold
            # Ustawienie czcionki dla znaków wschodnioazjatyckich
            run._element.rPr.rFonts.set(qn('w:eastAsia'), self.default_font)
        
        # Ustawienie odstępu po paragrafie
        paragraph.paragraph_format.space_after = self.spacing_after
    
    def create_document(self, words_text: str, title: str = "Lista słówek") -> BytesIO:
        """
        Tworzy dokument Word z listą słówek
        
        Args:
            words_text: Tekst z listą słówek wygenerowany przez AI
            title: Tytuł dokumentu (opcjonalnie)
            
        Returns:
            Obiekt BytesIO zawierający dokument Word
        """
        # Tworzenie nowego dokumentu Word
        doc = Document()
        
        # Ustawienie wąskich marginesów
        self._set_narrow_margins(doc)
        
        # Ustawienie dwóch kolumn
        self._set_two_columns(doc)
        
        # Przetwarzanie tekstu słówek linia po linii
        lines = words_text.strip().split('\n')
        
        # Wzorzec do wyciągania słówek
        # Format: numer. słówko (wymowa) – tłumaczenie
        word_pattern = r'^(\d+)\.\s+(.+?)\s*\(([^)]+)\)\s*[–-]\s*(.+)$'
        
        for line in lines:
            # Usuwanie białych znaków z początku i końca linii
            line = line.strip()
            
            # Pomijanie pustych linii - dodajemy pusty paragraf
            if not line:
                paragraph = doc.add_paragraph('')
                paragraph.paragraph_format.space_after = Pt(0)
                continue
            
            # Sprawdzanie czy linia to separator (linia z myślnikami)
            if re.match(r'^-+$', line):
                # Dodanie separatora jako tekst
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(line)
                run.font.name = self.default_font
                run.font.size = self.default_font_size
                paragraph.paragraph_format.space_after = self.spacing_after
                continue
            
            # Sprawdzanie czy linia to nagłówek kategorii (same wielkie litery)
            if line.isupper() and len(line) > 2:
                # Dodanie nagłówka kategorii - pogrubiony
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(line)
                run.font.name = self.default_font
                run.font.size = self.default_font_size
                run.bold = True
                paragraph.paragraph_format.space_after = self.spacing_after
                continue
            
            # Sprawdzanie czy to linia ze słówkiem
            word_match = re.match(word_pattern, line)
            if word_match:
                # Tworzenie paragrafu ze słówkiem
                paragraph = doc.add_paragraph()
                
                # Wyciąganie danych ze słówka
                number = word_match.group(1)
                english = word_match.group(2).strip()
                pronunciation = word_match.group(3).strip()
                polish = word_match.group(4).strip()
                
                # Numer i kropka - BEZ pogrubienia
                run_number = paragraph.add_run(f"{number}. ")
                run_number.font.name = self.default_font
                run_number.font.size = self.default_font_size
                run_number.bold = False
                
                # Słówko, wymowa i tłumaczenie - POGRUBIONE
                run_word = paragraph.add_run(f"{english} ({pronunciation}) – {polish}")
                run_word.font.name = self.default_font
                run_word.font.size = self.default_font_size
                run_word.bold = True
                
                # Ustawienie odstępu
                paragraph.paragraph_format.space_after = self.spacing_after
                continue
            
            # Sprawdzanie czy to linia z przykładem (ex: ...)
            if line.lower().startswith('ex:'):
                # Przykład - bez pogrubienia
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(line)
                run.font.name = self.default_font
                run.font.size = self.default_font_size
                run.bold = False
                paragraph.paragraph_format.space_after = self.spacing_after
                continue
            
            # Inna linia - bez pogrubienia
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(line)
            run.font.name = self.default_font
            run.font.size = self.default_font_size
            run.bold = False
            paragraph.paragraph_format.space_after = self.spacing_after
        
        # Zapisanie dokumentu do bufora pamięci
        buffer = BytesIO()
        doc.save(buffer)
        
        # Przewinięcie bufora na początek
        buffer.seek(0)
        
        return buffer
    
    def parse_words_from_text(self, text: str) -> list:
        """
        Parsuje tekst i wyciąga listę słówek
        
        Args:
            text: Tekst z listą słówek
            
        Returns:
            Lista słowników z informacjami o słówkach
        """
        # Lista na sparsowane słówka
        words = []
        
        # Wzorzec do wyciągania słówek
        # Format: numer. słówko (wymowa) – tłumaczenie
        pattern = r'(\d+)\.\s+([a-zA-Z\s]+)\s*\(([^)]+)\)\s*[–-]\s*(.+)'
        
        # Wzorzec do wyciągania przykładów
        example_pattern = r'ex:\s*(.+)'
        
        # Przetwarzanie tekstu linia po linii
        lines = text.strip().split('\n')
        current_word = None
        current_category = None
        
        for line in lines:
            line = line.strip()
            
            # Pomijanie pustych linii i separatorów
            if not line or re.match(r'^-+$', line):
                continue
            
            # Sprawdzanie czy to kategoria
            if line.isupper() and len(line) > 2:
                current_category = line
                continue
            
            # Próba dopasowania wzorca słówka
            word_match = re.match(pattern, line)
            if word_match:
                # Jeśli było poprzednie słówko, dodaj je do listy
                if current_word:
                    words.append(current_word)
                
                # Tworzenie nowego słówka
                current_word = {
                    'number': int(word_match.group(1)),        # Numer słówka
                    'english': word_match.group(2).strip(),    # Słówko angielskie
                    'pronunciation': word_match.group(3).strip(),  # Wymowa
                    'polish': word_match.group(4).strip(),     # Tłumaczenie polskie
                    'example': None,                           # Przykład (do uzupełnienia)
                    'category': current_category               # Kategoria
                }
                continue
            
            # Próba dopasowania przykładu
            example_match = re.match(example_pattern, line, re.IGNORECASE)
            if example_match and current_word:
                # Dodanie przykładu do aktualnego słówka
                current_word['example'] = example_match.group(1).strip()
        
        # Dodanie ostatniego słówka
        if current_word:
            words.append(current_word)
        
        return words
    
    def extract_word_list(self, text: str) -> list:
        """
        Wyciąga samą listę słówek angielskich (bez dodatkowych informacji)
        Używane do sprawdzania duplikatów
        
        Args:
            text: Tekst z listą słówek
            
        Returns:
            Lista słówek angielskich (stringów)
        """
        # Parsowanie pełnych słówek
        words = self.parse_words_from_text(text)
        
        # Wyciągnięcie samych słówek angielskich
        return [word['english'].lower() for word in words]
